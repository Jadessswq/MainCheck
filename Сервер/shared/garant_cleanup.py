"""
Очистка документов, выгруженных из «Гарант» и «КонсультантПлюс».

Убирает служебные элементы, которые засоряют контекст модели и путают её:
колонтитулы, маркеры системы, даты печати, копирайт, пустые строки,
обрезанные по ширине страницы разрывы строк, оглавление «по страницам».

Поддерживаются форматы: .txt, .rtf (best-effort), .doc/.docx (через python-docx),
.pdf (через pypdf — если установлен), .htm/.html (через BeautifulSoup — если установлен).
Для форматов без установленного парсера возвращается исходный текст как есть.
"""
from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List


# ── Регулярки служебных элементов ─────────────────────────────
_GARANT_HEADER = re.compile(
    r"^\s*(?:Система\s+ГАРАНТ|ГАРАНТ\s*[—–-]\s*правовая\s+поддержка).*$",
    re.IGNORECASE,
)
_KPLUS_HEADER = re.compile(
    r"^\s*КонсультантПлюс(?:\s+надежная\s+правовая\s+поддержка)?.*$",
    re.IGNORECASE,
)
_KPLUS_FOOTER = re.compile(
    r"^\s*(?:www\.consultant\.ru|КонсультантПлюс\s*\|\s*Дата).*$",
    re.IGNORECASE,
)
_PRINT_DATE = re.compile(
    r"^\s*Дата\s+(?:печати|сохранения|формирования)\s*[:\-]?\s*\d{1,2}[./-]\d{1,2}[./-]\d{2,4}.*$",
    re.IGNORECASE,
)
_COPY_MARK = re.compile(r"^\s*©\s*(?:(?:ООО|ЗАО|Компания)\s+)?«?НПП?\s*«?Гарант.*$", re.IGNORECASE)
_COPY_MARK_SIMPLE = re.compile(r"^\s*©\s*Гарант.*\d{4}.*$", re.IGNORECASE)
_COPY_MARK_KP = re.compile(r"^\s*©\s*КонсультантПлюс.*$", re.IGNORECASE)
_DOC_MARK = re.compile(r"^\s*Документ предоставлен\s+КонсультантПлюс.*$", re.IGNORECASE)
_PAGE_NUM = re.compile(r"^\s*[-–—]?\s*Страница\s+\d+\s+из\s+\d+\s*[-–—]?\s*$", re.IGNORECASE)
_PAGE_NUM_2 = re.compile(r"^\s*\d{1,4}\s*/\s*\d{1,4}\s*$")  # «12 / 48»
_DIV_LINE = re.compile(r"^\s*[-_=*─━═•·\.]{5,}\s*$")  # разделительные линии
_FORM_FEED = re.compile(r"[\x0c\x0b]")  # разрыв страницы
_MULTISPACE = re.compile(r"[ \t]{2,}")
_MULTI_NL = re.compile(r"\n{3,}")
# Мягкие переносы (посередине абзаца), когда слово разорвано: «слово-\nпродолжение»
_SOFT_HYPHEN = re.compile(r"([а-яёА-ЯЁa-zA-Z])-\n([а-яёa-z])")
# Жёсткие переносы внутри предложения без точки в конце
_HARD_WRAP = re.compile(r"([а-яёa-zA-Z,;:])\n(?=[а-яёa-z])")

# «Информация об изменениях», сноски, примечания ГАРАНТа — часто вставляются
# как отдельный абзац и ломают связность текста нормы.
_INFO_BLOCK = re.compile(
    r"^\s*(?:Информация об изменениях|Об ожидаемых изменениях|Внимание!).*$",
    re.IGNORECASE,
)
# «См. также», «См. комментарии» — гиперссылочные сноски
_SEE_ALSO = re.compile(r"^\s*См\.\s+(?:также|комментарии|Энциклопедии)\b.*$", re.IGNORECASE)

# Оглавление: строки вида «1.2.3. Название    123» (точка + пробелы + номер страницы)
_TOC_LINE = re.compile(r"^\s*[0-9IVX.]+\s+[^\n]+?[\s.]{3,}\d{1,4}\s*$")


_SKIP_PATTERNS: List[re.Pattern[str]] = [
    _GARANT_HEADER, _KPLUS_HEADER, _KPLUS_FOOTER, _PRINT_DATE,
    _COPY_MARK, _COPY_MARK_SIMPLE, _COPY_MARK_KP,
    _DOC_MARK, _PAGE_NUM, _PAGE_NUM_2, _DIV_LINE,
    _INFO_BLOCK, _SEE_ALSO, _TOC_LINE,
]


@dataclass
class CleanupStats:
    original_lines: int
    kept_lines: int
    removed_service: int
    joined_wrapped: int

    def as_dict(self) -> dict:
        return {
            "original_lines": self.original_lines,
            "kept_lines": self.kept_lines,
            "removed_service": self.removed_service,
            "joined_wrapped": self.joined_wrapped,
            "reduction_pct": round(
                100.0 * (1.0 - self.kept_lines / self.original_lines) if self.original_lines else 0.0,
                1,
            ),
        }


def _normalise_spaces(text: str) -> str:
    # Замена невидимых/специальных пробелов на обычный. Используем список,
    # чтобы не терять значимые символы (№ → «No» под NFKC, тонкий пробел и т.п.).
    replacements = {
        "\u00a0": " ",   # NBSP
        "\u202f": " ",   # narrow no-break space
        "\u2009": " ",   # thin space
        "\u2007": " ",   # figure space
        "\u2060": "",    # word joiner
        "\u200b": "",    # zero-width space
        "\ufeff": "",    # BOM
        "\u00ad": "",    # soft hyphen (невидимый — модель их путает)
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text


def clean_text(raw: str) -> tuple[str, CleanupStats]:
    """Очищает сырой текст, возвращает (текст, статистика)."""
    # Базовая канонизация без NFKC: NFKC разрушает «№» → «No», ломая
    # русские нормативные документы. Используем NFC + адресные замены.
    text = unicodedata.normalize("NFC", raw)
    text = _normalise_spaces(text)
    # Унификация переносов строк
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Убираем «разрыв страницы» и прочие управляющие
    text = _FORM_FEED.sub("\n", text)
    # Склеиваем разорванные переносом слова
    before = text
    text = _SOFT_HYPHEN.sub(r"\1\2", text)
    joined = before.count("-\n") - text.count("-\n")

    lines = text.split("\n")
    original_lines = len(lines)
    kept: List[str] = []
    removed = 0
    for line in lines:
        s = line.strip()
        skip = False
        for pat in _SKIP_PATTERNS:
            if pat.match(s):
                skip = True
                break
        if skip:
            removed += 1
            continue
        # Нормализуем внутренние пробелы
        s = _MULTISPACE.sub(" ", line)
        kept.append(s.rstrip())

    cleaned = "\n".join(kept)
    # Склеиваем «жёсткие переносы» внутри предложения: строчная после переноса
    before2 = cleaned.count("\n")
    cleaned = _HARD_WRAP.sub(r"\1 ", cleaned)
    joined += before2 - cleaned.count("\n")
    # Сжимаем пустые строки
    cleaned = _MULTI_NL.sub("\n\n", cleaned).strip()

    stats = CleanupStats(
        original_lines=original_lines,
        kept_lines=len(kept),
        removed_service=removed,
        joined_wrapped=max(joined, 0),
    )
    return cleaned, stats


# ── Извлечение текста из файла ────────────────────────────────
def _read_txt(path: Path) -> str:
    for enc in ("utf-8", "cp1251", "utf-16"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _read_docx(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise RuntimeError("Для .docx установите: pip install python-docx")
    doc = Document(str(path))
    parts: List[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        raise RuntimeError("Для .pdf установите: pip install pypdf")
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_html(path: Path) -> str:
    try:
        from bs4 import BeautifulSoup  # type: ignore
    except ImportError:
        # Примитивный fallback: стрипуем теги регуляркой
        data = _read_txt(path)
        return re.sub(r"<[^>]+>", " ", data)
    data = _read_txt(path)
    soup = BeautifulSoup(data, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    return soup.get_text("\n")


def _read_rtf(path: Path) -> str:
    # Базовая очистка RTF без внешних зависимостей.
    raw = _read_txt(path)
    # Удалить { \rtf … } командные токены
    no_cmds = re.sub(r"\\[a-zA-Z]+-?\d*\s?", "", raw)
    no_braces = re.sub(r"[{}]", "", no_cmds)
    # Декодирование \'XX последовательностей в cp1251
    def _repl(m: re.Match[str]) -> str:
        try:
            return bytes.fromhex(m.group(1)).decode("cp1251", errors="replace")
        except Exception:
            return ""
    return re.sub(r"\\'([0-9a-fA-F]{2})", _repl, no_braces)


_EXTRACTORS = {
    ".txt": _read_txt,
    ".md":  _read_txt,
    ".log": _read_txt,
    ".docx": _read_docx,
    ".pdf":  _read_pdf,
    ".htm":  _read_html,
    ".html": _read_html,
    ".rtf":  _read_rtf,
}


def extract_and_clean(path: Path) -> tuple[str, CleanupStats]:
    """Считывает файл → извлекает текст → чистит служебку Гарант/К+. """
    ext = path.suffix.lower()
    reader = _EXTRACTORS.get(ext)
    if reader is None:
        raise ValueError(f"Неподдерживаемый формат: {ext}. Поддерживаются: {sorted(_EXTRACTORS)}")
    raw = reader(path)
    return clean_text(raw)


def chunk_text(text: str, *, chunk_chars: int = 1200, overlap: int = 150) -> Iterable[str]:
    """Режет текст на перекрывающиеся чанки, стараясь разрезать по границам абзацев."""
    if chunk_chars <= 0:
        raise ValueError("chunk_chars должен быть > 0")
    if overlap >= chunk_chars:
        overlap = chunk_chars // 4

    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    buf: List[str] = []
    buf_len = 0
    for para in paragraphs:
        if buf_len + len(para) + 2 > chunk_chars and buf:
            chunk = "\n\n".join(buf).strip()
            if chunk:
                yield chunk
            # Начинаем новый буфер с перекрытием по последнему абзацу
            carry = buf[-1]
            if len(carry) > overlap:
                carry = carry[-overlap:]
            buf = [carry]
            buf_len = len(carry)
        buf.append(para)
        buf_len += len(para) + 2

    if buf:
        tail = "\n\n".join(buf).strip()
        if tail:
            yield tail
